"""Load PDFs / DOCX / TXT into chunked LangChain `Document` objects."""
from __future__ import annotations

import hashlib
import io
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from .config import settings


# Splitter tuned for biomedical prose: long sentences, lots of parentheticals,
# section headers like "Methods", "Results". We split on paragraphs first.
_SEPARATORS = ["\n\n", "\n", ". ", "? ", "! ", "; ", ", ", " ", ""]


def _splitter() -> RecursiveCharacterTextSplitter:
    return RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=_SEPARATORS,
        length_function=len,
    )


def _doc_id(source: str, content: str) -> str:
    return hashlib.sha1(f"{source}:{content[:64]}".encode("utf-8")).hexdigest()[:16]


def _read_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n\n".join(pages)


def _read_docx(data: bytes) -> str:
    import docx  # python-docx

    doc = docx.Document(io.BytesIO(data))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _read_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore")


def load_bytes(filename: str, data: bytes, *, extra_metadata: dict | None = None) -> list[Document]:
    """Read a file's bytes into chunked Documents with provenance metadata."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        text = _read_pdf(data)
    elif suffix == ".docx":
        text = _read_docx(data)
    elif suffix in {".txt", ".md"}:
        text = _read_txt(data)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    text = text.strip()
    if not text:
        return []

    chunks = _splitter().split_text(text)
    now = datetime.now(timezone.utc).isoformat()

    docs: list[Document] = []
    for i, chunk in enumerate(chunks):
        meta = {
            "source": filename,
            "chunk_index": i,
            "chunk_total": len(chunks),
            "ingested_at": now,
            "doc_id": _doc_id(filename, chunk),
        }
        if extra_metadata:
            meta.update(extra_metadata)
        docs.append(Document(page_content=chunk, metadata=meta))
    return docs


def load_path(path: Path, *, extra_metadata: dict | None = None) -> list[Document]:
    """Read a file from disk."""
    return load_bytes(path.name, path.read_bytes(), extra_metadata=extra_metadata)


def load_directory(directory: Path) -> list[Document]:
    """Load every supported file under a directory (non-recursive)."""
    docs: list[Document] = []
    for p in sorted(directory.iterdir()):
        if p.is_file() and p.suffix.lower() in {".pdf", ".docx", ".txt", ".md"}:
            docs.extend(load_path(p))
    return docs
